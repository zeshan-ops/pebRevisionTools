#!/usr/bin/env python3
"""Parse the prose FC4 notes (.docx) into per-area structured JSON.

The document is already organised to the syllabus:

    Heading 1  -> syllabus area, prefixed with its number ("12 - The Law ...")
    Heading 2  -> learning outcome, prefixed with its letter ("a)", "(b)")
    Heading 3/4-> subsection within an outcome, usually a named provision
    Normal / List Paragraph -> body

Styling is inconsistent in places: some outcomes are tagged Heading 3 rather
than Heading 2. So an outcome boundary is any heading whose text starts with a
letter in parens AND whose letter is the next expected one for that area -- the
same contiguity guard the exam parser uses, for the same reason (prose contains
"a)"-shaped strings).

Output: content/notes/fc4-notes.json, plus a coverage report against
content/syllabus/fc4.json.
"""
import json, re, sys
from pathlib import Path
import docx

ROOT = Path(__file__).resolve().parent.parent
SRC = ROOT / "content/notes/source/FC4 Notes (Abas).docx"
DEST = ROOT / "content/notes/fc4-notes.json"

AREA_RE = re.compile(r"^\s*(\d{1,2})\s*[-–—]")     # "12 - ", "9 – "
OUTCOME_RE = re.compile(r"^\s*\(?([a-z])\)")                  # "a)", "(b)"

def main():
    if not SRC.exists():
        sys.exit(f"missing {SRC.relative_to(ROOT)}")
    syl = json.loads((ROOT / "content/syllabus/fc4.json").read_text())
    by_num = {a["number"]: a for a in syl["areas"]}

    doc = docx.Document(str(SRC))
    areas, area, outcome = [], None, None

    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        style = p.style.name
        lvl = int(style[-1]) if style.startswith("Heading") and style[-1].isdigit() else None

        if lvl == 1:
            m = AREA_RE.match(text)
            num = int(m.group(1)) if m else None
            area = {"number": num, "docHeading": text,
                    "slug": by_num[num]["slug"] if num in by_num else None,
                    "intro": [], "outcomes": []}
            areas.append(area); outcome = None
            continue

        if area is None:
            continue

        # An outcome boundary: letter-prefixed heading, contiguous from 'a'.
        if lvl in (2, 3, 4):
            m = OUTCOME_RE.match(text)
            expected = chr(ord("a") + len(area["outcomes"]))
            if m and m.group(1) == expected:
                outcome = {"letter": m.group(1), "docHeading": text,
                           "body": [], "subsections": []}
                area["outcomes"].append(outcome)
                continue
            if outcome is not None:
                outcome["subsections"].append({"heading": text, "body": []})
                continue

        entry = {"text": text, "bullet": style == "List Paragraph"}
        if outcome is None:
            area["intro"].append(entry)
        elif outcome["subsections"]:
            outcome["subsections"][-1]["body"].append(entry)
        else:
            outcome["body"].append(entry)

    tables = []
    for t in doc.tables:
        tables.append([[c.text.strip() for c in r.cells] for r in t.rows])

    DEST.write_text(json.dumps({
        "generatedBy": "scripts/parse_prose_notes.py",
        "source": SRC.name,
        "areas": areas,
        "tables": tables,
    }, indent=2, ensure_ascii=False) + "\n")

    # ---- coverage report ----
    print(f"wrote {DEST.relative_to(ROOT)}")
    print(f"  areas parsed: {len(areas)}   tables: {len(tables)}")
    unmatched = [a["docHeading"][:50] for a in areas if a["slug"] is None]
    if unmatched:
        print(f"  !! headings not matched to a syllabus area: {unmatched}")

    print("\n  coverage — notes vs syllabus learning outcomes")
    print("  (application outcomes 'Apply … to a scenario' excluded: they are exam skills, not content)")
    total_missing = []
    for s in syl["areas"]:
        got = next((a for a in areas if a["number"] == s["number"]), None)
        want = [o for o in s["learningOutcomes"] if not o["isApplication"]]
        have = {o["letter"] for o in got["outcomes"]} if got else set()
        missing = [o["letter"] for o in want if o["letter"] not in have]
        words = sum(len(e["text"].split())
                    for a in ([got] if got else [])
                    for o in a["outcomes"]
                    for e in o["body"] + [b for s2 in o["subsections"] for b in s2["body"]])
        flag = f"  MISSING {','.join(missing)}" if missing else ""
        print(f"    {s['number']:>2}  {len(have)}/{len(want)} outcomes  {words:>5} words{flag}")
        total_missing += [(s["number"], m) for m in missing]
    print(f"\n  outcomes with no notes: {len(total_missing)}")

if __name__ == "__main__":
    main()
